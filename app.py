import streamlit as st
import os
import tempfile
from pathlib import Path
import uuid
import json
from io import BytesIO
from docx import Document # Added for DOCX export

# LangChain imports
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import RetrievalQA, LLMChain
from langchain.prompts import PromptTemplate

# For credentials (will use Streamlit secrets)
import google.generativeai as genai

st.set_page_config(page_title="StudyGenie", layout="wide")

# --- Initialize Session State --- 
# Ensure all necessary keys are initialized at the start
if 'api_key_configured' not in st.session_state:
    st.session_state.api_key_configured = False
if 'user_google_api_key' not in st.session_state:
    st.session_state.user_google_api_key = None
if 'vectordb' not in st.session_state:
    st.session_state.vectordb = None
if 'qa_chain' not in st.session_state:
    st.session_state.qa_chain = None
if 'chunks' not in st.session_state:
    st.session_state.chunks = []
if 'pdf_processed' not in st.session_state:
    st.session_state.pdf_processed = False
if 'pdf_filename' not in st.session_state:
    st.session_state.pdf_filename = None
if 'study_materials' not in st.session_state: # Initialize study materials state
    st.session_state.study_materials = None
if 'last_question' not in st.session_state:
    st.session_state.last_question = ""
if 'last_answer' not in st.session_state:
    st.session_state.last_answer = None
if 'last_sources' not in st.session_state:
    st.session_state.last_sources = []

# --- Configuration & Secrets --- 

# --- API Key Handling --- 
st.sidebar.subheader("Step 1: Configure API Key")
st.sidebar.info("Your Google AI API Key is required to power the AI features (Q&A, study material generation). It's used only for this session and not stored permanently.")
# Use a form for the API key input and submit button
with st.sidebar.form(key='api_key_form'):
    user_api_key_input = st.text_input(
        "Enter your Google API Key", 
        type="password", 
        key="user_google_api_key_input", 
        help="Get your key from Google AI Studio. Your key is used only for this session.",
        value=st.session_state.user_google_api_key or "" # Pre-fill if already submitted
    )
    submitted = st.form_submit_button("Submit Key")

    if submitted and user_api_key_input:
        st.session_state['user_google_api_key'] = user_api_key_input
        st.session_state['api_key_configured'] = False # Reset config status on new key submission
        st.sidebar.success("API Key submitted. Configuring...")
        # Trigger reconfiguration attempt below
    elif submitted and not user_api_key_input:
        st.sidebar.warning("Please enter an API Key.")
        st.session_state['api_key_configured'] = False # Ensure it's marked as not configured

# Determine which API key to use (ONLY user-provided)
GOOGLE_API_KEY = st.session_state.get('user_google_api_key', None)

# Validate and configure Google AI Client
api_configured_successfully = st.session_state.get('api_key_configured', False)

if GOOGLE_API_KEY and not api_configured_successfully:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        st.session_state['api_key_configured'] = True
        api_configured_successfully = True # Update local variable for immediate use
        st.sidebar.success("✅ Google AI Client Configured Successfully!")
        st.rerun() # Rerun to update UI state based on successful configuration
    except Exception as e:
        st.error(f"Failed to configure Google AI Client: {e}")
        st.session_state['api_key_configured'] = False
        api_configured_successfully = False
        # Don't stop the app, allow user to retry
elif not GOOGLE_API_KEY:
    st.warning("💡 Please enter your Google API Key in the sidebar and click 'Submit Key' to enable AI features.")
    # Ensure dependent features know the API is not ready
    st.session_state['api_key_configured'] = False
    api_configured_successfully = False
elif api_configured_successfully:
    st.sidebar.success("✅ Google AI Client Configured.") # Show confirmation if already done

# --- Constants --- 
FAISS_INDEX_PATH = "./faiss_index_streamlit" # Changed directory name for clarity

# --- Core Functions (Adapted from Notebook) ---

def load_pdf(pdf_path):
    """Load a PDF file and return a list of documents"""
    try:
        loader = PyPDFLoader(pdf_path)
        return loader.load()
    except Exception as e:
        st.error(f"Error loading PDF: {e}")
        return None

def process_documents(documents):
    """Process and split the documents into chunks"""
    if not documents:
        return []
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(documents)
        return chunks
    except Exception as e:
        st.error(f"Error processing documents: {e}")
        return []

def create_vectordb(chunks):
    """Create and save a FAISS vector database from the document chunks"""
    if not chunks:
        return None
    if not api_configured_successfully: # Check if API key is configured
        st.error("API Key not configured. Cannot create embeddings.")
        return None
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # Ensure the directory exists
        os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
        # Create a Vector store using FAISS
        vectordb = FAISS.from_documents(
            documents=chunks,
            embedding=embeddings
        )
        vectordb.save_local(folder_path=FAISS_INDEX_PATH) # Save the FAISS index
        st.success(f"Vector database created and saved to {FAISS_INDEX_PATH}")
        return vectordb
    except Exception as e:
        st.error(f"Error creating FAISS vector database: {e}")
        return None

def get_vectordb():
    """Load an existing FAISS vector database"""
    if not os.path.exists(FAISS_INDEX_PATH) or not os.path.exists(os.path.join(FAISS_INDEX_PATH, 'index.faiss')):
        # Don't show warning here, handle in UI logic
        # st.warning(f"FAISS index not found at {FAISS_INDEX_PATH}. Please upload a PDF first.")
        return None
    if not api_configured_successfully: # Check if API key is configured
        st.error("API Key not configured. Cannot load embeddings.")
        return None
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # Load the FAISS index
        # Allow dangerous deserialization is required for FAISS loading
        vectordb = FAISS.load_local(
            folder_path=FAISS_INDEX_PATH,
            embeddings=embeddings,
            allow_dangerous_deserialization=True
        )
        st.info(f"Loaded existing vector database from {FAISS_INDEX_PATH}")
        return vectordb
    except Exception as e:
        st.error(f"Error loading FAISS vector database: {e}")
        return None

def setup_qa_system(vectordb):
    """Setup a RAG-based Q&A system"""
    if not vectordb:
        return None
    if not api_configured_successfully: # Check if API key is configured
        st.error("API Key not configured. Cannot initialize LLM for Q&A.")
        return None
    try:
        # Using ChatGoogleGenerativeAI wrapper for consistency
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash", # Using flash model for potentially faster Q&A
            google_api_key=GOOGLE_API_KEY,
            temperature=0.2,
            top_p=0.95,
            max_output_tokens=2048,
            convert_system_message_to_human=True # Often needed for chat models in RAG
        )

        # Custom prompt template from the notebook
        template = """
        You are an expert tutor named StudyGenie specializing in explaining complex topics clearly and concisely based on the provided files. Your goal is to help a user understand a specific concept using ONLY the text provided from these documents.

        **Instructions:**

        1. Read the user's {question}.
        2. Carefully analyze the {context} provided below (this context is extracted from the uploaded documents).
        3. Generate a concise, clear, and accurate explanation of the {question} based *strictly* on the information within the {context}. Do not add external knowledge or information not present in the documents. Keep the explanation focused and to the point (2-4 sentences).
        4. Generate 2-3 distinct, relevant follow-up questions that a learner might ask to deepen their understanding of your explanation or the concept itself, based on the document content.
        5. Suggest 1-2 relevant search terms or types of external resources (e.g., 'search for tutorials on [topic]', 'look for articles explaining [concept]') that could help the user learn more about the topic based on the context.
        6. Format your entire response *strictly* as a single, valid JSON object. Do NOT include any text before or after the JSON object itself. The JSON object must have these exact keys: "explanation" (string), "follow_up_questions" (list of strings), and "potential_resources" (list of strings, where each string is a suggested search term or resource type).

        **Context:**
        Context:
        {context}

        **Concept:**
        Concept:
        {question}

        **JSON Output (MUST be only valid JSON):**
        """

        PROMPT = PromptTemplate(
            template=template,
            input_variables=["context", "question"]
        )

        # Create a RAG chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectordb.as_retriever(search_kwargs={"k": 5}),
            chain_type_kwargs={"prompt": PROMPT},
            return_source_documents=True
        )
        return qa_chain
    except Exception as e:
        st.error(f"Error setting up Q&A system: {e}")
        return None

def ask_question(qa_chain, concept_to_explain):
    """Ask a question to the QA system"""
    if not qa_chain:
        st.warning("QA system not initialized. Please upload a PDF first.")
        return None, None
    if not api_configured_successfully: # Check if API key is configured
        st.error("API Key not configured. Cannot ask questions.")
        return None, None
    try:
        # The RetrievalQA chain expects the query in the 'query' key by default
        result = qa_chain.invoke({"query": concept_to_explain})
        answer_json_str = result.get("result")
        source_docs = result.get("source_documents", [])

        # Attempt to parse the JSON response from the LLM
        try:
            # Clean potential markdown code fences
            if answer_json_str.startswith("```json\n"):
                answer_json_str = answer_json_str[7:]
            if answer_json_str.endswith("\n```"):
                answer_json_str = answer_json_str[:-4]
            answer_json_str = answer_json_str.strip()

            answer_data = json.loads(answer_json_str)
            # Validate expected keys
            if not all(k in answer_data for k in ["explanation", "follow_up_questions", "potential_resources"]):
                 st.warning("AI response is valid JSON but missing expected keys. Displaying raw response.")
                 st.code(answer_json_str)
                 # Provide a default structure but indicate the issue
                 answer_data = {
                    "explanation": answer_data.get("explanation", "Error: AI response missing 'explanation' key."),
                    "follow_up_questions": answer_data.get("follow_up_questions", ["Error: AI response missing 'follow_up_questions' key."]),
                    "potential_resources": answer_data.get("potential_resources", ["Error: AI response missing 'potential_resources' key."])
                 }

        except json.JSONDecodeError as e:
            st.error(f"Failed to parse the response from the AI as JSON. Error: {e}")
            st.warning("Raw AI response (which is not valid JSON):")
            st.code(answer_json_str)
            answer_data = {"explanation": f"Error: Could not parse AI response. The AI did not return valid JSON. Raw response logged above.", "follow_up_questions": [], "potential_resources": []}
        except Exception as e:
             st.error(f"An unexpected error occurred while processing the AI response: {e}")
             st.code(answer_json_str if answer_json_str else "No response content available.")
             answer_data = {"explanation": f"Error: An unexpected error occurred. See logs.", "follow_up_questions": [], "potential_resources": []}


        return answer_data, source_docs
    except Exception as e:
        st.error(f"Error asking question: {e}")
        return None, None

# --- Study Material Generation Functions ---

def create_llm_for_generation():
    """Creates an LLM instance specifically for generation tasks."""
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro", # Using Pro model for potentially higher quality generation
            google_api_key=GOOGLE_API_KEY,
            temperature=0.3, # Slightly higher temp for more creative generation
            max_output_tokens=4096, # Allow longer outputs for summaries/MCQs
            convert_system_message_to_human=True
        )
        return llm
    except Exception as e:
        st.error(f"Error creating LLM for generation: {e}")
        return None

def create_summary_generator(llm):
    """Create a summary generator chain."""
    summary_template = """
    You are an expert educational summarizer. Create a clear, concise summary of the following text.
    Focus on the main concepts, key points, and critical details.

    TEXT:
    {text}

    SUMMARY:
    """
    summary_prompt = PromptTemplate(template=summary_template, input_variables=["text"])
    summary_chain = LLMChain(llm=llm, prompt=summary_prompt)
    return summary_chain

def create_mcq_generator(llm):
    """Create a multiple-choice question generator chain (outputting plain text)."""
    mcq_template = """
    You are an expert educator designing exam-applicable multiple-choice questions. Based *only* on the core concepts, key principles, and significant information presented in the following text, create {num_questions} multiple-choice questions.

    **Focus exclusively on understanding and application of the material.** Do NOT ask questions about:
    - Document metadata (e.g., page numbers, specific section titles unless core to the content)
    - The source of the information (e.g., "According to page 5...")
    - Trivial details or overly specific examples unless they illustrate a fundamental concept.

    Each question must test conceptual understanding and have 4 options (labeled A, B, C, D) with only one clearly correct answer. Format your response as plain text, following this structure EXACTLY for each question:

    Q[question number]: [Question text focusing on core concepts]
    A. [Option A]
    B. [Option B]
    C. [Option C]
    D. [Option D]
    Correct Answer: [Correct option letter, e.g., C]
    Explanation: [Brief explanation of why the answer is correct, linking back to the core concept]

    (Ensure there is a blank line between each question block)

    TEXT:
    {text}

    EXAM-APPLICABLE MULTIPLE CHOICE QUESTIONS FOCUSED ON CORE CONCEPTS:
    """

    mcq_prompt = PromptTemplate(
        template=mcq_template,
        input_variables=["text", "num_questions"]
    )
    mcq_chain = LLMChain(llm=llm, prompt=mcq_prompt)

    return mcq_chain

def create_flashcard_generator(llm):
    """Create a flashcard generator chain (outputting plain text)."""
    flashcard_template = """
    You are an expert educator. Create {num_cards} flashcards focusing on the CORE CONCEPTS, KEY TERMS, and MAIN IDEAS presented in the following text.
    Avoid trivial details or overly specific examples unless they are central to understanding a core concept.
    Format your response as PLAIN TEXT, following this structure EXACTLY for each card:

    Front: [Question, term, or concept]
    Back: [Answer, definition, or explanation]

    (Ensure there is a blank line between each card block)

    TEXT:
    {text}

    PLAIN TEXT FLASHCARDS (focus on core concepts):
    """
    flashcard_prompt = PromptTemplate(template=flashcard_template, input_variables=["text", "num_cards"])
    flashcard_chain = LLMChain(llm=llm, prompt=flashcard_prompt)
    return flashcard_chain

def generate_study_materials(chunks, num_chunks=5, num_mcqs=5, num_flashcards=5):
    """Generate study materials from selected chunks."""
    if not chunks:
        st.warning("No document chunks available to generate study materials.")
        return None

    llm = create_llm_for_generation()
    if not llm:
        return None

    # Create generator chains
    summary_generator = create_summary_generator(llm)
    mcq_generator = create_mcq_generator(llm)
    flashcard_generator = create_flashcard_generator(llm)

    # Select chunks (use more chunks for better context)
    selected_chunks = chunks[:min(num_chunks, len(chunks))]
    if not selected_chunks:
        st.warning("Not enough document chunks to generate materials.")
        return None
    combined_text = "\n\n".join([chunk.page_content for chunk in selected_chunks])

    results = {}
    try:
        st.write("Generating summary...")
        summary = summary_generator.invoke({"text": combined_text}).get("text", "Failed to generate summary.")
        results["summary"] = summary
    except Exception as e:
        st.error(f"Error generating summary: {e}")
        results["summary"] = "Error generating summary."

    try:
        st.write("Generating MCQs (as text)...")
        mcq_response = mcq_generator.invoke({"text": combined_text, "num_questions": num_mcqs}).get("text")
        # Store the raw text response directly, no JSON parsing needed
        results["mcqs_text"] = mcq_response if mcq_response else "MCQ generation failed or returned empty."
    except Exception as e:
        st.error(f"Error generating MCQs: {e}")
        results["mcqs_text"] = f"Error during MCQ generation: {e}"

    try:
        st.write("Generating flashcards (as text)...")
        flashcard_response = flashcard_generator.invoke({"text": combined_text, "num_cards": num_flashcards}).get("text")
        # Store the raw text response directly
        results["flashcards_text"] = flashcard_response if flashcard_response else "Flashcard generation failed or returned empty."
    except Exception as e:
        st.error(f"Error generating flashcards: {e}")
        results["flashcards_text"] = f"Error during flashcard generation: {e}"

    return results

# --- Helper functions for output formatting ---
def parse_flashcards_text(text):
    """Parses plain text flashcards into a list of dictionaries."""
    cards = []
    current_card = {}
    lines = text.strip().split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith("Front:"):
            # Save previous card if it was valid and complete
            if current_card and current_card.get("front") and current_card.get("back"):
                cards.append(current_card)
            current_card = {"front": line[len("Front:"):].strip(), "back": ""}
        elif line.startswith("Back:") and current_card:
            current_card["back"] = line[len("Back:"):].strip()
        # No need to check for blank lines explicitly, the logic handles it

    # Add the last card if it's valid and complete
    if current_card and current_card.get("front") and current_card.get("back"):
        cards.append(current_card)

    # Handle cases where the last line might be a 'Front:' without a 'Back:'
    # or if the input text is malformed. This ensures robustness.
    if not cards and current_card.get("front") and not current_card.get("back"):
         # If only a front was found for the very first/only card attempt
         st.warning("Found a flashcard front without a corresponding back. Check AI output format.")
         # Optionally add a placeholder or skip
         # cards.append({"front": current_card["front"], "back": "[Missing Back]"})

    elif not cards and text.strip(): # If no cards parsed but text exists
        st.warning("Could not parse any flashcards from the generated text. Check AI output format.")
        st.text_area("Raw Flashcard Text:", text, height=150)

    return cards


def create_docx(mcq_text):
    """Creates a DOCX file from the MCQ text."""
    document = Document()
    document.add_heading('StudyGenie MCQs', 0)
    document.add_paragraph(mcq_text)
    
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio

def create_markdown(mcq_text):
    """Creates Markdown content from the MCQ text."""
    # The text is already mostly markdown-like, just return it
    # Could add a header if desired
    markdown_content = f"# StudyGenie MCQs\n\n{mcq_text}"
    return markdown_content.encode('utf-8') # Encode to bytes for download

# --- Streamlit UI --- 
st.title("📚 StudyGenie: Your AI-Powered Study Assistant")
st.markdown("Upload your PDF document, ask questions, and generate study materials!")

# --- PDF Upload Section --- 
st.sidebar.subheader("Step 2: Upload Your PDF")
uploaded_file = st.sidebar.file_uploader("Choose a PDF file", type="pdf", key="pdf_uploader")

# --- PDF Processing Logic --- 
if uploaded_file is not None and api_configured_successfully:
    # Check if it's a new file or the same one
    new_filename = uploaded_file.name
    if new_filename != st.session_state.get('pdf_filename'):
        st.session_state.pdf_processed = False # Mark as not processed if filename changes
        st.session_state.vectordb = None
        st.session_state.qa_chain = None
        st.session_state.chunks = []
        st.session_state.study_materials = None # Reset study materials on new PDF
        st.session_state.pdf_filename = new_filename # Store the new filename
        st.session_state.last_question = ""
        st.session_state.last_answer = None
        st.session_state.last_sources = []
        st.info(f"New PDF '{new_filename}' detected. Ready for processing.")

    if not st.session_state.pdf_processed:
        with st.spinner(f"Processing '{new_filename}'..."):
            try:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name

                st.write("Loading PDF...")
                documents = load_pdf(tmp_file_path)

                if documents:
                    st.write("Splitting document into chunks...")
                    chunks = process_documents(documents)
                    st.session_state.chunks = chunks # Store chunks

                    if chunks:
                        st.write("Creating vector database (this might take a moment)...")
                        vectordb = create_vectordb(chunks)
                        st.session_state.vectordb = vectordb

                        if vectordb:
                            st.write("Setting up Q&A system...")
                            qa_chain = setup_qa_system(vectordb)
                            st.session_state.qa_chain = qa_chain
                            st.session_state.pdf_processed = True
                            st.success(f"✅ PDF '{new_filename}' processed successfully!")
                            st.rerun() # Rerun to update UI state
                        else:
                            st.error("Failed to create vector database.")
                    else:
                        st.error("Failed to process document chunks.")
                else:
                    st.error("Failed to load PDF document.")

            except Exception as e:
                st.error(f"An error occurred during PDF processing: {e}")
            finally:
                # Clean up temporary file
                if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                    os.remove(tmp_file_path)
    else:
        st.sidebar.success(f"✅ PDF '{st.session_state.pdf_filename}' already processed.")

elif uploaded_file is None:
    st.info("Upload a PDF in the sidebar to get started.")
    # Reset state if no file is uploaded
    if st.session_state.pdf_filename is not None:
        st.session_state.pdf_processed = False
        st.session_state.vectordb = None
        st.session_state.qa_chain = None
        st.session_state.chunks = []
        st.session_state.study_materials = None
        st.session_state.pdf_filename = None
        st.session_state.last_question = ""
        st.session_state.last_answer = None
        st.session_state.last_sources = []
        st.experimental_rerun()

elif not api_configured_successfully:
    st.warning("Please configure your API key in the sidebar to process PDFs.")


# --- Main Area Layout (Tabs) --- 
tab1, tab2 = st.tabs(["❓ Q&A", "📝 Study Materials"])

with tab1:
    st.header("Ask Questions About Your Document")

    if st.session_state.get('pdf_processed') and st.session_state.get('qa_chain') and api_configured_successfully:
        st.info("Enter a concept or question based on your uploaded PDF.")

        # Use a form for the question input
        with st.form(key='qa_form'):
            user_question = st.text_input(
                "Concept to Explain:",
                key="user_question_input",
                placeholder="e.g., What is the main idea of chapter 3?",
                value=st.session_state.get('last_question', '') # Pre-fill with last question
            )
            submit_question = st.form_submit_button("Get Explanation")

        if submit_question and user_question:
            st.session_state.last_question = user_question # Store the submitted question
            with st.spinner("StudyGenie is thinking..."):
                answer_data, source_docs = ask_question(st.session_state.qa_chain, user_question)
                st.session_state.last_answer = answer_data
                st.session_state.last_sources = source_docs
                # Rerun to display results outside the form scope immediately
                st.rerun()
        elif submit_question and not user_question:
            st.warning("Please enter a question or concept.")

        # Display the last answer if available
        if st.session_state.get('last_answer'):
            answer_data = st.session_state.last_answer
            source_docs = st.session_state.last_sources

            st.subheader(f"Explanation for: " + st.session_state.get('last_question', 'your query'))
            st.markdown(answer_data.get("explanation", "No explanation provided."))

            st.subheader("Follow-up Questions")
            follow_ups = answer_data.get("follow_up_questions", [])
            if follow_ups:
                for i, q in enumerate(follow_ups):
                    st.markdown(f"- {q}")
            else:
                st.markdown("_No follow-up questions suggested._")

            st.subheader("Potential Resources")
            resources = answer_data.get("potential_resources", [])
            if resources:
                for i, r in enumerate(resources):
                    st.markdown(f"- {r}")
            else:
                st.markdown("_No potential resources suggested._")

            # Display source documents (optional)
            if source_docs:
                with st.expander("View Sources from Document"):
                    for i, doc in enumerate(source_docs):
                        st.markdown(f"**Source {i+1} (Page {doc.metadata.get('page', 'N/A')})**")
                        st.caption(doc.page_content[:500] + "...") # Show snippet

    elif not st.session_state.get('pdf_processed'):
        st.info("Please upload and process a PDF first.")
    elif not api_configured_successfully:
        st.warning("API Key not configured. Please configure it in the sidebar.")
    else:
        st.warning("QA system not ready. Ensure PDF is processed and API key is valid.")

with tab2:
    st.header("Generate Study Materials")

    if st.session_state.get('pdf_processed') and st.session_state.get('chunks') and api_configured_successfully:

        # User controls for generation
        col1, col2 = st.columns(2)
        with col1:
            num_mcqs_input = st.number_input(
                "Number of Multiple Choice Questions:",
                min_value=1,
                max_value=50, # Set max limit
                value=5, # Default value
                step=1,
                key="num_mcqs"
            )
        with col2:
            num_flashcards_input = st.number_input(
                "Number of Flashcards:",
                min_value=1,
                max_value=50, # Set max limit
                value=5, # Default value
                step=1,
                key="num_flashcards"
            )

        if st.button("✨ Generate Study Materials", key="generate_materials_button"):
            with st.spinner("Generating study materials... This may take a while."):
                # Use user-selected values
                materials = generate_study_materials(
                    st.session_state.chunks,
                    num_chunks=10, # Use more chunks for better context
                    num_mcqs=num_mcqs_input,
                    num_flashcards=num_flashcards_input
                )
                st.session_state.study_materials = materials
                st.rerun() # Rerun to display generated materials

        # Display generated materials if they exist
        if st.session_state.get('study_materials'):
            materials = st.session_state.study_materials

            st.subheader("Generated Summary")
            st.markdown(materials.get("summary", "_Summary not generated._"))

            st.subheader("Generated Multiple Choice Questions (MCQs)")
            mcq_text = materials.get("mcqs_text", "_MCQs not generated._")
            st.text_area("MCQs", mcq_text, height=300, key="mcq_display_area")
            # Add download button for MCQs
            if mcq_text and not mcq_text.startswith("Error") and not mcq_text.startswith("MCQ generation failed"):
                try:
                    docx_bytes = create_docx(mcq_text)
                    st.download_button(
                        label="Download MCQs as DOCX",
                        data=docx_bytes,
                        file_name=f"{st.session_state.pdf_filename}_mcqs.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Failed to create DOCX for download: {e}")

            st.subheader("Generated Flashcards")
            flashcards_text = materials.get("flashcards_text", "_Flashcards not generated._")
            if flashcards_text and not flashcards_text.startswith("Error") and not flashcards_text.startswith("Flashcard generation failed"):
                parsed_cards = parse_flashcards_text(flashcards_text)
                if parsed_cards:
                    for i, card in enumerate(parsed_cards):
                        with st.expander(f"**Card {i+1}:** {card['front']}", expanded=False):
                            st.markdown(f"**Back:** {card['back']}")
                else:
                    st.warning("Could not parse flashcards from the generated text. Displaying raw output:")
                    st.text_area("Raw Flashcard Text", flashcards_text, height=200, key="flashcard_raw_display")
            else:
                st.markdown(flashcards_text) # Display error or 'not generated' message

    elif not st.session_state.get('pdf_processed'):
        st.info("Please upload and process a PDF first to enable study material generation.")
    elif not api_configured_successfully:
        st.warning("API Key not configured. Please configure it in the sidebar.")
    else:
        st.warning("Document chunks not available. Ensure PDF was processed correctly.")

# --- Footer or other UI elements ---
st.sidebar.markdown("--- ")
st.sidebar.info("StudyGenie uses Google Generative AI.")

# --- Attempt to load existing VectorDB on startup if PDF was processed before ---
# This logic might need refinement depending on desired persistence
# if not st.session_state.get('vectordb') and os.path.exists(FAISS_INDEX_PATH) and api_configured_successfully:
#     with st.spinner("Loading existing knowledge base..."):
#         st.session_state.vectordb = get_vectordb()
#         if st.session_state.vectordb:
#             st.session_state.qa_chain = setup_qa_system(st.session_state.vectordb)
#             # Assume chunks are needed if DB exists? This is tricky without saving chunks.
#             # Maybe prompt user to re-upload if chunks aren't in session state?
#             st.sidebar.success("Loaded existing knowledge base.")
#             st.session_state.pdf_processed = True # Assume processed if DB loaded
#             # What filename to associate? Needs better state management for persistence.
#             # st.session_state.pdf_filename = "previously_processed.pdf" # Placeholder
#         else:
#             st.sidebar.warning("Could not load existing knowledge base.")